from flask import Flask, render_template, request, send_file, jsonify
from docx import Document
from docx.enum.section import WD_ORIENT
from docx.shared import Inches
from docx.shared import Pt
import re
import datetime
from datetime import date
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from openpyxl import Workbook
from openpyxl.drawing.image import Image
from openpyxl.styles import PatternFill, Border, Side, Alignment, Font, NamedStyle

app = Flask(__name__)

@app.route('/')
def main_page():
    return render_template('main.html')

wb = Workbook()
ws = wb.active
# Значения переменных для примера
date = '01/01/2024'
name_requestor = 'Requestor Name'
name_owner = 'Owner Name'
text = 'Description text'
link = 'http://linktoapproval.com'
day_of_week = 'Monday'
path = '/path/to/template'
time = '12:00 PM'
# Определение стилей
header_style = NamedStyle(name="headerStyle",
                          fill=PatternFill("solid", fgColor="FFFF00"),
                          font=Font(bold=True),
                          border=Border(left=Side(style='thin'),
                                        right=Side(style='thin'),
                                        top=Side(style='thin'),
                                        bottom=Side(style='thin')),
                          alignment=Alignment(horizontal="center", vertical="center"))

data_style = NamedStyle(name="dataStyle",
                        border=Border(left=Side(style='thin'),
                                      right=Side(style='thin'),
                                      top=Side(style='thin'),
                                      bottom=Side(style='thin')),
                        alignment=Alignment(horizontal="left", vertical="center"))

# Применение стилей
wb.add_named_style(header_style)
wb.add_named_style(data_style)

data = [
    ('Date', '23.03.2024'),
    ('CRM call or CBCP Project name *', ''),
    ('Change Requestor (link or description to request)', 'Dmitriy Yershov / Tair Babenov'),
    ('Change Owner (IT Representative)', 'Alexandr Kondratyev / Sergei Kotelnikov')
]

# Определяем индекс, где должен быть вставлен 'Change definition'
change_definition_index = 4

# Заполнение данных до 'Change definition'
for row_index, (title, content) in enumerate(data, start=1):
    title_cell = f'A{row_index}'
    content_cell = f'B{row_index}'
    ws[title_cell] = title
    ws[title_cell].style = header_style
    ws[content_cell] = content
    ws[content_cell].style = data_style
    ws.row_dimensions[row_index].height = 40
    ws.merge_cells(start_row=row_index, start_column=2, end_row=row_index, end_column=4)

# Вставка 'Change definition'
ws.insert_rows(change_definition_index)
change_definition_cell = f'A{change_definition_index}'
ws.merge_cells(f'A{change_definition_index}:D{change_definition_index}')
ws[change_definition_cell] = 'Change definition'
ws[change_definition_cell].style = header_style
ws.row_dimensions[change_definition_index].height = 40

# После 'Change definition' добавляем оставшиеся данные
data2 = [
    ('Change Type', 'Standard'),
    ('Change Priority', 'Medium'),
    ('Change description / Justification for change',
     'Deployment of April 2024 Microsoft security updates...'),
]

# Заполнение оставшихся данных
for row_index, (title, content) in enumerate(data2, start=change_definition_index + 1):
    title_cell = f'A{row_index}'
    content_cell = f'B{row_index}'
    ws[title_cell] = title
    ws[title_cell].style = header_style
    ws[content_cell] = content
    ws[content_cell].style = data_style
    ws.row_dimensions[row_index].height = 40
    ws.merge_cells(start_row=row_index, start_column=2, end_row=row_index, end_column=4)
# Установка ширины столбцов
ws.column_dimensions['A'].width = 50
ws.column_dimensions['B'].width = 50
ws.column_dimensions['C'].width = 20
ws.column_dimensions['D'].width = 50

# Сохранение файла
file_path = './Formatted_Change_Control_Form.xlsx'
wb.save(file_path)

@app.route('/generate_excel_deployment', methods=["POST"])
def download_excel_deployment():
    # Получаем ввод пользователя из запроса
    user_choices = request.form.to_dict()

    wb = Workbook()
    ws = wb.active

    # Устанавливаем название рабочего листа
    ws.title = "Change Control Form"

    # Определяем стили
    header_fill = PatternFill(start_color="FFFF00", end_color="FFFF00", fill_type="solid")
    header_font = Font(bold=True)
    center_aligned_text = Alignment(horizontal="center")
    thin_border = Border(left=Side(style='thin'), right=Side(style='thin'),
                         top=Side(style='thin'), bottom=Side(style='thin'))

    # Устанавливаем ширину столбцов
    ws.column_dimensions['A'].width = 20
    ws.column_dimensions['B'].width = 20
    ws.column_dimensions['C'].width = 20
    ws.column_dimensions['D'].width = 20

    # Добавляем изображение
    img = Image('/image/img.png')  # Убедитесь, что путь к изображению корректен
    ws.add_image(img, 'A1:D2')

    # Применяем стили к заголовку и подзаголовку
    ws.merge_cells('A3:D3')
    ws['A3'] = 'IT&T Change Control Form'
    ws['A3'].fill = header_fill
    ws['A3'].font = header_font
    ws['A3'].alignment = center_aligned_text
    ws['A3'].border = thin_border

    # Добавляем секцию выбора
    ws['A4'] = 'PLEASE SELECT THE CHANGE:'
    ws['A4'].font = Font(bold=True)

    # Устанавливаем высоту строки для чекбоксов
    ws.row_dimensions[5].height = 30

    # Вставляем заглушки для чекбоксов и текст
    for col, option in enumerate(["GENERAL", "SERVICE ACCOUNT"], start=1):
        cell = ws.cell(row=5, column=col)
        cell.value = option
        cell.border = thin_border
        cell.alignment = center_aligned_text
        # Добавляем галочку, если пользователь сделал выбор
        if user_choices.get(option) == 'on':
            ws.cell(row=5, column=col).value = f"\u2713 {option}"

    data = {
        ('Date', f'{date}'),
        ('CRM call or CBCP Project name *', ''),
        ('Change Requestor (link or description to request)', f'{name_requestor}'),
        ('Change Owner (IT Representative)', f'{name_owner}'),
    }
    ws['A10'] = 'Change defenition'
    data2 = {
        ('Change Type', 'Standard'),
        ('Change Priority', 'Medium'),
        ('Change description / \nJustification for change', f'{text}'),
        ('', ''),
        ('Data Owner Approval (link)', f'{link}')
    }
    ws['A16'] = 'Change documents'
    ws['A17'] = 'Risk Assessments'
    data3 = {
        ('', 'Probability', 'Impact', 'Risk Level'),
        ('Risk of implementing the change and its impacts to production', 'Unlikely', 'Moderate', 'Medium')
    }
    data4 = {
        ('Test plan (link or description)', 'Test Plan'),
        ('Detailed Change implementation plan (link or description)', 'Implementation Plan'),
        ('Roll back plan (link or description)', 'Rollback Plan'),
        ('User message (link), if required', f'Notification message #1\nTo be sent on {day_of_week} {date} to users of Test Group #2 computers (these PCs are members of "WSUS_WKS_TEST" AD Group).\nPath to the message template: {path}\n\n Notification message #2\nTo be sent on {day_of_week} {date} to members of distribution list "AKS - All KPO Users".\nPath to the message template: {path}'),
    }
    ws['A24'] = ('Implementation date', f'{date}', 'Time', f'{time}')

    data5 = {
        ('Total outage (including contingency time)', 'About 15 minutes'),
        ('List of affected systems (hardware)', 'Laptops, desktops and workstations'),
        ('Application(s) affected (software)','Windows 11 Operating System, Windows 10 Operating System, Microsoft Office 2021, Microsoft Office 2019, Microsoft Office 2016, Microsoft Office 2013'),
        ('Groups of Users affected', 'All KPO users except the users of computers that are members of AD Group "SUS Exclusion Group".'),
        ('MyApp Portal Data update required ', 'N'),
        ('Resources required to implement change (manpower)', 'IT Infrastructure Engineer: Downloading updates and creation of software update group.\bIT&T Service Desk: Sending the user notification messages.\bIT Support Engineer: Creation of the target collections, arranging MS updates distribution, monitoring and control the deployment process.'),
        ('Third party / supplier involvement *', 'N')
    }
    ws['A35'] = 'Business Continuity / Disaster Recovery Planning checklist'

    data6 = {
        ('Require IT BCP server update (Y/N)', 'N'),
        ('RI (recovery instruction) update (Y/N)', 'N'),
    }

    ws['A42'] = 'IT&T Service Desk use only'
    ws['A43'] = ('CRM work order number for RFC', '')
    ws['A44'] = 'Post Change Analysis'

    data7 = {
        ('Actual downtime', ''),
        ('Post installation checks', ''),
        ('Problems arising from change', ''),
        ('MyApp Portal updated', ''),
        ('Date')
    }
    ws['A50'] = 'Note: fields marked * are optional                                                                       Revised: 10/12/2020'

    # Сохраняем рабочую книгу
    file_path = 'Change_Control_Form_User_Choices.xlsx'
    wb.save(file_path)

    # Возвращаем файл для скачивания
    return send_file(file_path, as_attachment=True)

def clean_xml_string(value):
    """Remove characters not compatible with XML."""
    # This regex will match all characters that are not ASCII or printable, excluding valid whitespace characters.
    return re.sub(r'[^\x20-\x7E]+', '', value)

@app.route('/generate_word_plan', methods=['POST'])
def download_word_plan():
    doc = Document()
    section = doc.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE

    section.page_width = Inches(11.69)
    section.page_height = Inches(8.27)

    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(9)

    date2 = clean_xml_string(request.form['date'])
    today = date.today()
    # Форматируем дату в строку (например, в формате ГГГГ-ММ-ДД)
    date_string = today.strftime('%Y-%m-%d')
    name = clean_xml_string(request.form['name'])
    speciality = clean_xml_string(request.form['speciality'])
    path = clean_xml_string(request.form['path'])

    # Преобразование строки даты в объект datetime
    date_obj = datetime.datetime.strptime(date2, '%Y-%m-%d')
    month = date_obj.month

    # Determine the parity of the month
    parity = "odd" if month % 2 != 0 else "even"
    # Получение дня недели
    day_of_week = date_obj.strftime('%A')
    title_table = doc.add_table(rows=2, cols=4)
    title_table.style = 'Table Grid'

    # Add a table to the Word document
    table = doc.add_table(rows=1, cols=3)
    table.style = 'Table Grid'

    title_cells = title_table.rows[0].cells
    title_cells[0].text = 'Title'
    title_cells[1].text = 'Test Plan'
    title_cells[2].text = 'VSM Call No.'
    title_cells[3].text = ''  # Assuming you fill this programmatically or leave it empty

    creator_cells = title_table.rows[1].cells
    creator_cells[0].text = 'Created by'
    creator_cells[1].text = name + ' - ' + speciality
    creator_cells[2].text = 'Date'
    creator_cells[3].text = date_string  # Insert the date from the form

    # Format the title table
    for row in title_table.rows:
        for cell in row.cells:
            cell.paragraphs[0].runs[0].font.bold = True  # Make text bold
            shading_elm = OxmlElement('w:shd')
            shading_elm.set(qn('w:fill'), 'D3D3D3')  # Grey background
            cell._tc.get_or_add_tcPr().append(shading_elm)

    # Define the headers
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Step No.'
    hdr_cells[1].text = 'Description'
    hdr_cells[2].text = 'Additional conditions / Details'

    # Format the header row
    for cell in hdr_cells:
        shading_elm = OxmlElement('w:shd')
        shading_elm.set(qn('w:fill'), 'D3D3D3')
        cell._tc.get_or_add_tcPr().append(shading_elm)

    # Add other rows of the table
    steps_data = [
        ('1', 'Test Phase – Deployment', ''),
        ('1.1', f'Create the Software Update Group “WKS MS updates released in {parity} month” in SCCM.',
         'Path to Software Update Group in Configuration Manager Console: \\Software Library\\Overview\\Software Updates\\Software Update Groups'),
        ('1.2',
         'Deploy the Software Update Group “WKS MS updates released in even month” to target collection “[TEST1] WKS MS updates - RestrictReboot”.',
         f'Software available time: {date2} 16:00\nInstallation deadline: {date2} 17:00.\n\nMaintenance Windows: Yes\nSuppress the system restart: Yes\nOverride Maintenance Window: Yes'),
        ('1.3',
         'Deploy the Software Update Group “WKS MS updates released in even month” to target collection “[TEST1] WKS MS updates - Restart”.\nPath to target collection in Configuration Manager Console: \\Assets and Compliance\\Overview\\Device Collections\\KPO Custom Collections\\MS Updates Folder\\Workstations\\1_Test',
         f'Software available time: {date2} 16:00.\nInstallation deadline: {date2} 17:00.\n\nMaintenance Windows: Yes\nSuppress the system restart: No\nOverride Maintenance Window: Yes'),
        ('1.4', 'Check results of the MS updates deployment to Test Group #1 computers.', ''),
        ('2', 'Pre-Production Phase – Notification', ''),
        ('2.1', 'Group #2 computers which are members of AD Group “WSUS_WKS_TEST”.Send notification message “Notification_Message_1.msg” to users of Test ',
         f'Date: {day_of_week} {date2}.\n\nPath to Notification Message #1:“{path}”.\n'),
        ('3', 'Pre-Production Phase – Deployment', ''),
        ('3.1', 'Deploy the Software Update Group “WKS MS updates released in even month” to target collection [TEST2] WKS MS updates - RestrictReboot”.\n\nPath to target collections in Configuration Manager Console:\\Assets and Compliance\\Overview\\Device Collections\\KPO Custom Collections\\MS Updates Folder\\Workstations\\2_PreProduction\n\n',
         f'Software available time: {date2} 16:00\nInstallation deadline: {date2} 17:00.\n\nMaintenance Windows: Yes\nSuppress the system restart: Yes\nOverride Maintenance Window: Yes'),
        ('3.2', 'Deploy the Software Update Group “WKS MS updates released in even month” to target collection “[TEST2] WKS MS updates - Restart”.',
         f'Software available time: {date2} 16:00\nInstallation deadline: {date2} 17:00.\n\nMaintenance Windows: Yes\nSuppress the system restart: Yes\nOverride Maintenance Window: Yes'),
        ('3.3','Deploy the Software Update Group “WKS MS updates released in even month” to target collection “[TEST2-MP14] WKS MS updates - RestrictReboot”.',
         f'Software available time: {date2} 16:00\nInstallation deadline: {date2} 17:00.\n\nMaintenance Windows: Yes\nSuppress the system restart: Yes\nOverride Maintenance Window: Yes'),
        ('3.4', 'Deploy the Software Update Group “WKS MS updates released in even month” to target collection “[TEST2-MP14] WKS MS updates - Restart”.',
         f'Software available time: {date2} 16:00\nInstallation deadline: {date2} 17:00.\n\nMaintenance Windows: Yes\nSuppress the system restart: Yes\nOverride Maintenance Window: Yes'),
        ('4', 'Control the Deployment', ''),
        ('4.1', 'Control the deployment status by one of the following ways: \nWith use of Monitoring Workspace in Configuration Manager Console; \nWith use of SCCM reports http://ccmrep/', ''),
        ('5', 'Check the Deployment Results', ''),
        ('5.1', 'Check the test computers and analyse the deployment results. \n If deployment is successful, proceed with deployment to computers in production environment.', '')

    ]

    for step_no, description, details in steps_data:
        row_cells = table.add_row().cells
        row_cells[0].text = step_no
        row_cells[1].text = description
        row_cells[2].text = details

    # Save the document
    filename = 'Plan2.docx'
    doc.save(filename)

    # Serve the document as a download
    return send_file(filename, as_attachment=True)

@app.route('/generate_word_implementation', methods=['POST'])
def download_word_implementation():
    doc = Document()
    section = doc.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE

    section.page_width = Inches(11.69)
    section.page_height = Inches(8.27)

    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(9)

    date2 = clean_xml_string(request.form['date'])
    today = date.today()
    # Форматируем дату в строку (например, в формате ГГГГ-ММ-ДД)
    date_string = today.strftime('%Y-%m-%d')
    name = clean_xml_string(request.form['name'])
    speciality = clean_xml_string(request.form['speciality'])
    path = clean_xml_string(request.form['path'])

    # Преобразование строки даты в объект datetime
    date_obj = datetime.datetime.strptime(date2, '%Y-%m-%d')
    month = date_obj.month

    # Determine the parity of the month
    parity = "odd" if month % 2 != 0 else "even"
    # Получение дня недели
    day_of_week = date_obj.strftime('%A')


    title_table = doc.add_table(rows=2, cols=4)
    title_table.style = 'Table Grid'

    # Add a table to the Word document
    table = doc.add_table(rows=1, cols=3)
    table.style = 'Table Grid'

    title_cells = title_table.rows[0].cells
    title_cells[0].text = 'Title'
    title_cells[1].text = 'Implementation Plan'
    title_cells[2].text = 'VSM Call No.'
    title_cells[3].text = ''  # Assuming you fill this programmatically or leave it empty

    creator_cells = title_table.rows[1].cells
    creator_cells[0].text = 'Created by'
    creator_cells[1].text = name + ' - ' + speciality
    creator_cells[2].text = 'Date'
    creator_cells[3].text = date_string  # Insert the date from the form

    # Format the title table
    for row in title_table.rows:
        for cell in row.cells:
            cell.paragraphs[0].runs[0].font.bold = True  # Make text bold
            shading_elm = OxmlElement('w:shd')
            shading_elm.set(qn('w:fill'), 'D3D3D3')  # Grey background
            cell._tc.get_or_add_tcPr().append(shading_elm)

    # Define the headers
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Step No.'
    hdr_cells[1].text = 'Description'
    hdr_cells[2].text = 'Additional conditions / Details'

    # Format the header row
    for cell in hdr_cells:
        shading_elm = OxmlElement('w:shd')
        shading_elm.set(qn('w:fill'), 'D3D3D3')
        cell._tc.get_or_add_tcPr().append(shading_elm)

    # Add other rows of the table
    steps_data = [
        ('1', 'Production Phase – Notification', ''),
        ('1.1', 'Send notification message “Notification_Message_2.msg” to members of distribution list "AKS - All KPO Users".',
         f'Date: {day_of_week} {date2} \n\n Path to Notification Message #2: {path}'),
        ('2', 'Production Phase – Deployment', ''),
        ('2.1',
         f'Deploy the Software Update Group “WKS MS updates released in {parity} month” to target collection “[PROD] WKS MS updates - RestrictReboot”.',
         f'Software available time: {date2} 12:00.\nInstallation deadline: {date2} 22:00.\n\nMaintenance Windows: Yes\nSuppress the system restart: No\nOverride Maintenance Window: Yes\n\nPath to target collections in Configuration Manager Console:\\Assets and Compliance\\Overview\\Device Collections\\KPO Custom Collections\\MS Updates Folder\\Workstations\\2_PreProduction'),
        ('2.2', f'Deploy the Software Update Group “WKS MS updates released in {parity} month” to target collection “[PROD] WKS MS updates - Restart”.',
         f'Software available time: {date2} 16:00\nInstallation deadline: {date2} 22:10.\n\nMaintenance Windows: Yes\nSuppress the system restart: Yes\nOverride Maintenance Window: Yes'),
        ('2.3', f'Deploy the Software Update Group “WKS MS updates released in {parity} month” to target collection “[PROD-MP14] WKS MS updates - RestrictReboot”.',
         f'Software available time: {date2} 16:00\nInstallation deadline: {date2} 22:00.\n\nMaintenance Windows: Yes\nSuppress the system restart: Yes\nOverride Maintenance Window: Yes'),
        ('2.4',
         f'Deploy the Software Update Group “WKS MS updates released in {parity} month” to target collection “[PROD] WKS MS updates - Restart”.',
         f'Software available time: {date2} 16:00\nInstallation deadline: {date2} 22:10.\n\nMaintenance Windows: Yes\nSuppress the system restart: Yes\nOverride Maintenance Window: Yes'),
        ('3', 'Control the Deployment', ''),
        ('3.1', 'Control the deployment status by one of the following ways: \nWith use of Monitoring Workspace in Configuration Manager Console; \nWith use of SCCM reports http://ccmrep/', ''),
        ('4', 'Check the Deployment Results', ''),
        ('5.1', 'Check test computers and analyse the deployment results. If deployment is successful, close the IT Request for Change.', '')

    ]

    for step_no, description, details in steps_data:
        row_cells = table.add_row().cells
        row_cells[0].text = step_no
        row_cells[1].text = description
        row_cells[2].text = details

    # Save the document
    filename = 'Implementation.docx'
    doc.save(filename)

    # Serve the document as a download
    return send_file(filename, as_attachment=True)

@app.route('/generate_word_rollback', methods=['POST'])
def download_word_rollback():
    doc = Document()
    section = doc.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    spacer_paragraph = doc.add_paragraph()
    spacer_run = spacer_paragraph.add_run()
    spacer_run.add_break()

    # Set the spacing after the paragraph, which effectively creates space between tables
    # This example sets 12 points of spacing, but you can adjust as necessary
    paragraph_format = spacer_paragraph.paragraph_format
    paragraph_format.space_after = Pt(12)
    section.page_width = Inches(11.69)
    section.page_height = Inches(8.27)

    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(9)

    today = date.today()
    # Форматируем дату в строку (например, в формате ГГГГ-ММ-ДД)
    date_string = today.strftime('%Y-%m-%d')
    name = clean_xml_string(request.form['name'])
    speciality = clean_xml_string(request.form['speciality'])

    title_table = doc.add_table(rows=2, cols=4)
    title_table.style = 'Table Grid'

    # Add a table to the Word document
    table = doc.add_table(rows=1, cols=3)
    table.style = 'Table Grid'

    title_cells = title_table.rows[0].cells
    title_cells[0].text = 'Title'
    title_cells[1].text = 'Rollback Plan'
    title_cells[2].text = 'VSM Call No.'
    title_cells[3].text = ''  # Assuming you fill this programmatically or leave it empty

    creator_cells = title_table.rows[1].cells
    creator_cells[0].text = 'Created by'
    creator_cells[1].text = name + ' - ' + speciality
    creator_cells[2].text = 'Date'
    creator_cells[3].text = date_string  # Insert the date from the form

    # Format the title table
    for row in title_table.rows:
        for cell in row.cells:
            cell.paragraphs[0].runs[0].font.bold = True  # Make text bold
            shading_elm = OxmlElement('w:shd')
            shading_elm.set(qn('w:fill'), 'D3D3D3')  # Grey background
            cell._tc.get_or_add_tcPr().append(shading_elm)

    # Define the headers
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Step No.'
    hdr_cells[1].text = 'Description'
    hdr_cells[2].text = 'Additional conditions / Details'

    # Format the header row
    for cell in hdr_cells:
        shading_elm = OxmlElement('w:shd')
        shading_elm.set(qn('w:fill'), 'D3D3D3')
        cell._tc.get_or_add_tcPr().append(shading_elm)

    # Add other rows of the table
    steps_data = [
        ('1', 'Create the package to uninstall failed updates on affected PCs.', '')
    ]

    for step_no, description, details in steps_data:
        row_cells = table.add_row().cells
        row_cells[0].text = step_no
        row_cells[1].text = description
        row_cells[2].text = details

        desc_paragraph = row_cells[1].paragraphs[0]
        desc_paragraph.add_run(description).bold = True

    # Save the document
    filename = 'Rollback.docx'
    doc.save(filename)

    # Serve the document as a download
    return send_file(filename, as_attachment=True)

@app.route('/implementation')
def implementation():
    return render_template('implementation.html')

@app.route('/rollback')
def rollback():
    return render_template('rollback.html')

@app.route('/excel')
def test():
    return render_template('deployment.html')


if __name__ == '__main__':
    app.run(debug=True)